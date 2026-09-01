from langgraph.graph import StateGraph, START, END
from typing import TypedDict, Annotated,  Optional
from langchain_core.messages import  BaseMessage, HumanMessage, SystemMessage,AIMessage
from langchain_ollama import ChatOllama
from langchain_groq import ChatGroq
from dotenv import load_dotenv
# from langgraph.checkpoint.memory import MemorySaver
from langgraph.checkpoint.sqlite import SqliteSaver

from langgraph.graph.message import add_messages
import os
from langchain_openai import ChatOpenAI
import sqlite3
import uuid


#  tools--
from langgraph.prebuilt import ToolNode, tools_condition
from langchain_tavily import TavilySearch
from langchain_core.tools import tool
from langchain_core.documents import Document
# -------rag based system--------------
from langchain_chroma import Chroma
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
# ----------------------imports----------
# ---------------------hitl imports-----------
from langgraph.types import interrupt
from langgraph.config import get_stream_writer
# -------------------------------------------

import requests
import math

load_dotenv()

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
STOCK_API_KEY = os.getenv("STOCK_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")


def groq_model():
    """Groq's hosted API -- fast, and the only backend here that actually
    works on a cloud host like Render (no local model daemon required).
    This is the PRIMARY model: fast enough that even simple prompts don't
    look 'stuck', unlike the local Ollama model."""
    if not GROQ_API_KEY:
        print("GROQ_API_KEY not set -- skipping Groq.")
        return None
    try:
        llm = ChatGroq(
            # small/cheap model -- lower token cost than larger models,
            # at some quality tradeoff. (llama-3.1-8b-instant is no longer
            # available on Groq; this is the current smallest general
            # chat-capable model with tool-calling support.)
            model="openai/gpt-oss-20b",
            api_key=GROQ_API_KEY,
            temperature=0.3,
        )
        llm.invoke("ping")  # confirm the key/model actually work, not just that the object constructed
        return llm
    except Exception as e:
        print("Could not connect to Groq:", e)
    return None


def laptop_model():
    """Local Ollama -- only works on a machine that has Ollama running.
    Kept as a LOCAL-DEV-ONLY fallback for offline use; never reachable on
    Render or any other host without an Ollama daemon."""
    try:
        llm = ChatOllama(model="qwen2.5:3b", temperature=0.3)
        # a real call is needed to confirm Ollama is actually reachable —
        # constructing ChatOllama never fails on its own, even if Ollama
        # isn't running
        llm.invoke("ping")
        return llm
    except Exception as e:
        print("Could not connect to Ollama. Make sure Ollama app is running on your machine:", e)
    return None


def online_model():
    """Optional third option: a custom OpenAI-compatible endpoint, if
    LOCAL_MODEL_URL/API_KEY/NAME are configured. Not used unless the other
    two aren't available."""
    try:
        llm = ChatOpenAI(
            base_url=os.environ["LOCAL_MODEL_URL"] + "/v1",
            api_key=os.environ["LOCAL_MODEL_API_KEY"],
            model=os.environ["LOCAL_MODEL_NAME"],
            temperature=0.3,
            # the phone's cache_proxy.py buffers the full response and
            # always returns plain JSON, never real SSE — requesting a
            # stream here gets back application/json instead of
            # text/event-stream, which the client's SSE parser can't
            # read, raising "No generations found in stream"
            streaming=False,
        )
        return llm
    except Exception as e:
        print("Could not connect to OpenAI. Make sure OpenAI app is running on your machine:", e)
    return None


# Groq first -- fast and works on Render/any cloud host. Falls back to the
# local Ollama model only if Groq isn't configured/reachable (e.g. running
# fully offline during local dev without a GROQ_API_KEY).
llm = groq_model()

# Exposed to the frontend so it can show which backend is actually active
# (and whether it's the slow local-fallback path) via a status indicator,
# instead of a generic spinner with no explanation.
ACTIVE_MODEL_INFO = {
    "backend": "groq",
    "model_name": "openai/gpt-oss-20b",
    "is_slow_fallback": False,
}

if llm is None:
    llm = laptop_model()
    ACTIVE_MODEL_INFO = {
        "backend": "ollama",
        "model_name": "qwen2.5:3b",
        "is_slow_fallback": True,
    }


if llm == None:
    print("No model is available. Closing...")
    exit(1)
# -------------------embeddings-----------
# Lazily constructed: FastEmbedEmbeddings pulls in onnxruntime and downloads/
# loads the model weights on first use, which is a real memory spike. On
# Render's free 512MB tier, doing that at import time means the process is
# already carrying that weight before a single request even arrives, leaving
# less headroom for the actual embedding work during a file upload (which is
# what pushed memory to ~530MB and got the process OOM-killed). Deferring
# construction to the first actual call means idle memory stays lower, and
# the model loads only once (cached) after that.
_embedding_model = None


def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = FastEmbedEmbeddings(model_name='BAAI/bge-small-en-v1.5')
    return _embedding_model


def is_embedding_model_loaded():
    """Lets the frontend show a one-time 'loading the embedding model,
    this can take ~1 minute on Render's free tier' explanation only on the
    FIRST upload/RAG use in a process's lifetime, instead of every time."""
    return _embedding_model is not None


# --------------------- chunk storage split: SQLite + Chroma ---------------
# Chroma always stores whatever text you hand it alongside each vector (its
# similarity_search returns that text directly) -- there's no "vectors only"
# mode. To keep Chroma's own index smaller, we embed and store only a
# SHORT/truncated form of each chunk there (enough for the embedding model
# to place it well, and for a human-readable preview), while the FULL chunk
# text lives in a separate SQLite table. Retrieval finds matching chunk IDs
# via Chroma similarity search, then loads the full text for those IDs from
# SQLite -- keeping Chroma's per-chunk payload small without losing any
# content.
_CHUNK_DB_PATH = "./resume_chunks.db"
_CHROMA_PREVIEW_CHARS = 200  # how much of each chunk Chroma itself stores/embeds against


def _get_chunk_db():
    """Lazily opens the chunk-text SQLite connection and ensures the table
    exists. Deferred (not opened at import time) for the same reason the
    embedding model is lazy: keep baseline memory/handles low until a file
    is actually uploaded."""
    conn = sqlite3.connect(_CHUNK_DB_PATH, check_same_thread=False)
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS chunks (
            chunk_id TEXT PRIMARY KEY,
            source TEXT,
            page INTEGER,
            content TEXT,
            scope TEXT NOT NULL DEFAULT 'temporary'
        )
        """
    )
    # in case an older DB from before `scope` existed is still around
    cols = [row[1] for row in conn.execute("PRAGMA table_info(chunks)").fetchall()]
    if "scope" not in cols:
        conn.execute("ALTER TABLE chunks ADD COLUMN scope TEXT NOT NULL DEFAULT 'temporary'")
    conn.commit()
    return conn


def store_chunks(chunks, scope="temporary", id_prefix=""):
    """Persists the full text of each chunk in SQLite, keyed by a stable
    chunk_id, and returns the list of chunk_ids in the same order -- used
    both as the SQLite primary key and as the Chroma document id, so the
    two stores can be joined by id.

    `scope` is either "temporary" (the current resume upload -- cleared out
    whenever a new resume is uploaded, keeping the single shared collection
    from growing unbounded) or "permanent" (info the user explicitly asked
    to remember via remember_info_tool, kept across uploads).
    `id_prefix` disambiguates chunk_ids between temporary/permanent content
    that might otherwise collide (e.g. same source path re-ingested)."""
    conn = _get_chunk_db()
    chunk_ids = []
    rows = []
    for index, chunk in enumerate(chunks):
        source = chunk.metadata.get("source", "")
        page = chunk.metadata.get("page", index)
        chunk_id = f"{id_prefix}{scope}::{source}::{page}::{index}"
        chunk_ids.append(chunk_id)
        rows.append((chunk_id, source, page, chunk.page_content, scope))

    conn.executemany(
        "INSERT OR REPLACE INTO chunks (chunk_id, source, page, content, scope) VALUES (?, ?, ?, ?, ?)",
        rows,
    )
    conn.commit()
    conn.close()
    return chunk_ids


def delete_temporary_chunks():
    """Removes all TEMPORARY chunks (the previously uploaded resume) from
    both SQLite and the shared Chroma collection -- called right before
    ingesting a newly uploaded resume, so the single collection holds only
    the current resume's temporary data plus whatever permanent info has
    been explicitly remembered, instead of accumulating every past upload
    forever."""
    conn = _get_chunk_db()
    rows = conn.execute("SELECT chunk_id FROM chunks WHERE scope = 'temporary'").fetchall()
    temp_ids = [row[0] for row in rows]
    conn.execute("DELETE FROM chunks WHERE scope = 'temporary'")
    conn.commit()
    conn.close()

    if temp_ids:
        vector_store = Chroma(
            embedding_function=get_embedding_model(),
            collection_name='my_pdf_document',
            persist_directory='./my_chroma_db',
        )
        vector_store.delete(ids=temp_ids)

    return len(temp_ids)


def load_chunk_texts(chunk_ids):
    """Fetches full chunk text for the given ids from SQLite, returned as a
    {chunk_id: text} dict."""
    if not chunk_ids:
        return {}
    conn = _get_chunk_db()
    placeholders = ",".join("?" for _ in chunk_ids)
    rows = conn.execute(
        f"SELECT chunk_id, content FROM chunks WHERE chunk_id IN ({placeholders})",
        chunk_ids,
    ).fetchall()
    conn.close()
    return dict(rows)


# Tracks the most recently uploaded resume / job-description file paths --
# used by ats_score_tool and resume_review_tool, which need the FULL
# document text rather than top-k similarity chunks.
_last_uploaded_pdf_path = None
_last_uploaded_jd_path = None


def _load_docx_text(file_path):
    """Extract plain text from a .docx file using python-docx directly
    (avoids depending on the docx2txt package that langchain's
    Docx2txtLoader needs but isn't installed here)."""
    import docx
    document = docx.Document(file_path)
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_full_text(file_path):
    """Extract the full plain text of a resume/JD file -- supports PDF,
    DOCX, and plain TXT, the formats accepted by the chat file uploader."""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        docs = PyPDFLoader(file_path).load()
        return "\n\n".join(doc.page_content for doc in docs)
    if lower.endswith(".docx"):
        return _load_docx_text(file_path)
    if lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    raise ValueError(f"Unsupported file type: {file_path}")


# --------------------- dynamic chunker (metaprogramming registry) ---------
# A decorator-based strategy registry: each chunking strategy is a plain
# function registered under a name via @register_chunker(...). The actual
# strategy used for a given upload is picked at runtime by
# pick_chunker_strategy(), based on the document's size -- so a 2-page note
# and a 1000-page notebook don't get chunked with the same fixed
# chunk_size/overlap. New strategies (e.g. a markdown-aware or
# heading-aware splitter) can be added later just by writing a function and
# decorating it; nothing else in the ingestion pipeline needs to change.

_CHUNKER_REGISTRY = {}


def register_chunker(name):
    """Decorator: registers a chunking-strategy function under `name` in
    the module-level strategy registry, so pick_chunker_strategy() can look
    it up by name at runtime instead of hardcoding which splitter to use."""
    def decorator(fn):
        _CHUNKER_REGISTRY[name] = fn
        return fn
    return decorator


@register_chunker("small")
def _chunk_small(docs):
    # short documents (a few pages): small chunks, generous overlap, for
    # precise retrieval when there isn't much text to search through anyway
    return RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=150)


@register_chunker("medium")
def _chunk_medium(docs):
    # the original default -- typical reports/resumes/papers, tens of pages
    return RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)


@register_chunker("large")
def _chunk_large(docs):
    # long documents (roughly 100-500 pages): bigger chunks keep the total
    # chunk count (and embedding cost) from scaling linearly with page
    # count, while a smaller overlap ratio still preserves context at chunk
    # boundaries without excessive duplication
    return RecursiveCharacterTextSplitter(chunk_size=1800, chunk_overlap=250)


@register_chunker("huge")
def _chunk_huge(docs):
    # very large documents -- e.g. a ~1000-page notebook/manual. Chunk size
    # scales further so the number of chunks (and thus embedding calls and
    # Chroma index size) stays manageable; overlap stays modest since with
    # this many chunks a large overlap ratio would balloon storage for
    # little retrieval benefit
    return RecursiveCharacterTextSplitter(chunk_size=2500, chunk_overlap=300)


def pick_chunker_strategy(docs):
    """Chooses a registered chunking strategy based on the document's size
    (page count, falling back to total character count for single-page
    text extractions like .docx/.txt). Returns the strategy name and the
    RecursiveCharacterTextSplitter instance to use."""
    page_count = len(docs)
    total_chars = sum(len(doc.page_content) for doc in docs)

    if page_count >= 700 or total_chars >= 1_500_000:
        name = "huge"
    elif page_count >= 100 or total_chars >= 250_000:
        name = "large"
    elif page_count >= 15 or total_chars >= 30_000:
        name = "medium"
    else:
        name = "small"

    return name, _CHUNKER_REGISTRY[name](docs)


# --------------------- dynamic HITL step registry (metaprogramming) -------
# Same pattern as the chunker registry above: each HITL checkpoint is a
# plain handler function registered under a name via @register_hitl_step.
# A tool that needs human approval doesn't call interrupt() and hand-roll
# its own decision logic -- it calls run_hitl_step(name, question, context)
# and the registry dispatches to the right handler at runtime. A handler
# can return a `next_step` name, letting run_hitl_step chain straight into
# the next registered checkpoint (e.g. approve the ATS statement -> then
# ask about a docx/pdf export) without the calling tool hardcoding that
# sequence. New steps are added later purely by writing + decorating a new
# handler -- nothing about run_hitl_step or existing steps needs to change.

_HITL_STEP_REGISTRY = {}


def register_hitl_step(name):
    """Decorator: registers a HITL step handler under `name`. The handler
    receives (decision, context) and returns a result dict; optionally
    including "next_step" to chain into another registered step."""
    def decorator(fn):
        _HITL_STEP_REGISTRY[name] = fn
        return fn
    return decorator


def run_hitl_step(name, question, context=None, extra_payload=None):
    """Pauses the graph with interrupt(), tagging the payload with which
    step this is, then dispatches the human's decision to the handler
    registered under `name`. If that handler's result includes a
    "next_step", automatically continues into that next registered step
    (chaining checkpoints dynamically instead of the caller hardcoding a
    fixed sequence of interrupt() calls)."""
    context = context or {}
    payload = {"question": question, "step": name}
    if extra_payload:
        payload.update(extra_payload)

    decision = interrupt(payload)

    handler = _HITL_STEP_REGISTRY.get(name)
    if handler is None:
        return {"status": "error", "message": f"No HITL handler registered for step {name!r}"}

    result = handler(decision, context)

    next_step = result.get("next_step")
    if next_step:
        return run_hitl_step(
            next_step,
            result["next_question"],
            context={**context, **result.get("next_context", {})},
        )

    return result


@register_hitl_step("ats_statement")
def _hitl_ats_statement(decision, context):
    approved = isinstance(decision, str) and decision.lower() in ("yes", "approved")
    statement = context.get("suggested_statement", "")

    result = {
        "status": "approved" if approved else "rejected",
        "message": (
            f"The user APPROVED adding this ATS-boosting statement -- include it "
            f"(filled in with real details from the resume/JD, not the bracket "
            f"placeholders) in an appropriate section: “{statement}”"
            if approved else
            "The user REJECTED adding the suggested ATS-boosting statement -- "
            "do NOT add it. Proceed with the rest of the tailoring only."
        ),
    }

    # dynamically chain into the export-format checkpoint next, rather than
    # the calling tool hardcoding "after this step, ask about export"
    if context.get("offer_export"):
        result["next_step"] = "offer_file_export"
        result["next_question"] = (
            "Your tailored resume text is ready. Would you like me to also "
            "prepare it as a downloadable file? Reply with 'docx', 'pdf', or "
            "'no' (a plain .txt/.md download is always available either way)."
        )

    return result


@register_hitl_step("offer_file_export")
def _hitl_offer_file_export(decision, context):
    choice = decision.strip().lower() if isinstance(decision, str) else ""

    if choice in ("docx", "word"):
        export_format = "docx"
    elif choice == "pdf":
        export_format = "pdf"
    else:
        export_format = None

    return {
        "status": "export_choice",
        "export_format": export_format,
        "message": (
            f"The user requested a {export_format.upper()} export -- mention "
            f"that a {export_format} download will be prepared in addition to "
            f"the .txt/.md version."
            if export_format else
            "The user did not request a docx/pdf export -- only the .txt/.md "
            "download applies."
        ),
    }


# --------------------- dynamic progress/ETA reporting ---------------------
# Metaprogramming: each tool's pipeline is registered as an ORDERED
# sequence of named stages via @register_progress_stages(tool_name), one
# decorator call per tool instead of hand-writing status strings inline
# wherever a tool happens to be invoked. Adding progress reporting for a
# new tool later is just one decorated list -- nothing about the
# frontend's rendering loop changes.
#
# DSA: each stage sequence is walked like a simple FIFO queue (a Python
# list used front-to-back), and get_stage_progress() computes a PREFIX SUM
# over each stage's estimated_seconds to answer "how much time has
# elapsed by the time we reach stage i" and "how much total ETA remains"
# in O(1) per lookup after an O(n) prefix pass -- the same prefix-sum
# technique used for range-sum queries, applied here to time estimates
# instead of array values.

_PROGRESS_STAGE_REGISTRY = {}


def register_progress_stages(tool_name):
    """Decorator: registers the ordered list of {"name", "estimated_seconds"}
    stage descriptors returned by the decorated function, under `tool_name`,
    looked up at runtime by get_stage_progress() to report
    done/doing/next/ETA to the user."""
    def decorator(fn):
        _PROGRESS_STAGE_REGISTRY[tool_name] = fn()
        return fn
    return decorator


@register_progress_stages("ats_score_tool")
def _ats_score_tool_stages():
    return [
        {"name": "Loading uploaded resume", "estimated_seconds": 1},
        {"name": "Retrieving relevant resume sections", "estimated_seconds": 2},
        {"name": "Scoring ATS compatibility with the LLM", "estimated_seconds": 20},
    ]


@register_progress_stages("resume_review_tool")
def _resume_review_tool_stages():
    return [
        {"name": "Loading uploaded resume", "estimated_seconds": 1},
        {"name": "Retrieving relevant resume sections", "estimated_seconds": 2},
        {"name": "Waiting for your ATS-statement approval", "estimated_seconds": 0},
        {"name": "Waiting for your export-format choice", "estimated_seconds": 0},
        {"name": "Generating tailored resume with the LLM", "estimated_seconds": 25},
    ]


@register_progress_stages("ats_tailor_pipeline_tool")
def _ats_tailor_pipeline_tool_stages():
    return [
        {"name": "Loading uploaded resume", "estimated_seconds": 1},
        {"name": "Loading/finding the job description", "estimated_seconds": 3},
        {"name": "Scoring current ATS compatibility", "estimated_seconds": 15},
        {"name": "Finding missing keywords and weaknesses", "estimated_seconds": 15},
        {"name": "Waiting for your ATS-statement approval", "estimated_seconds": 0},
        {"name": "Waiting for your export-format choice", "estimated_seconds": 0},
        {"name": "Generating the tailored, ATS-optimized resume", "estimated_seconds": 25},
    ]


@register_progress_stages("search_tool")
def _search_tool_stages():
    return [
        {"name": "Searching the web (Tavily)", "estimated_seconds": 4},
    ]


@register_progress_stages("rag_tool")
def _rag_tool_stages():
    return [
        {"name": "Retrieving relevant document chunks", "estimated_seconds": 2},
    ]


def get_stage_progress(tool_name, stage_index):
    """Given a tool name and the index of the stage currently running,
    returns what's done, what's doing, what's next, and the ETA (seconds)
    to finish the remaining stages -- via a prefix-sum pass over the
    registered stage list. Returns None if the tool has no registered
    stages (caller should fall back to a generic status message)."""
    stages = _PROGRESS_STAGE_REGISTRY.get(tool_name)
    if not stages:
        return None
    stage_index = max(0, min(stage_index, len(stages) - 1))

    # prefix sum: cumulative[i] = total estimated seconds for stages[0..i]
    cumulative = []
    running_total = 0
    for stage in stages:
        running_total += stage["estimated_seconds"]
        cumulative.append(running_total)

    total_eta = cumulative[-1]
    elapsed_estimate = cumulative[stage_index] - stages[stage_index]["estimated_seconds"]
    remaining_eta = total_eta - elapsed_estimate

    return {
        "done": [s["name"] for s in stages[:stage_index]],
        "doing": stages[stage_index]["name"],
        "next": stages[stage_index + 1]["name"] if stage_index + 1 < len(stages) else None,
        "eta_seconds_remaining": remaining_eta,
        "total_stages": len(stages),
        "current_stage_number": stage_index + 1,
    }


def format_progress_label(tool_name, stage_index):
    """Builds a single human-readable status line ('done: X | doing: Y
    (~Ns) | next: Z') for the given tool/stage, or a generic fallback if
    the tool has no registered progress stages."""
    progress = get_stage_progress(tool_name, stage_index)
    if progress is None:
        return f"🔎 Running `{tool_name}`..."

    parts = [f"⏳ Step {progress['current_stage_number']}/{progress['total_stages']}: {progress['doing']}"]
    if progress["done"]:
        parts.append(f"(done: {', '.join(progress['done'])})")
    if progress["eta_seconds_remaining"]:
        parts.append(f"~{progress['eta_seconds_remaining']}s remaining")
    if progress["next"]:
        parts.append(f"next: {progress['next']}")
    return " | ".join(parts)


def report_progress(tool_name, stage_index):
    """Called from inside a tool to emit a live progress update for the
    given stage, via LangGraph's stream writer (stream_mode="custom").
    Safe to call even outside a streaming graph run (e.g. direct unit
    tests / tool.invoke() calls) -- get_stream_writer() raises either
    RuntimeError ("outside of a runnable context") or KeyError
    ('__pregel_runtime' missing, e.g. when a tool is .invoke()'d directly
    rather than run through chatbot.stream()) in that case, both swallowed."""
    try:
        writer = get_stream_writer()
    except (RuntimeError, KeyError):
        return
    writer({
        "progress_update": {
            "tool": tool_name,
            "stage_index": stage_index,
            "label": format_progress_label(tool_name, stage_index),
        }
    })


# -----------rag function for pdf file traverse-------------
def ingest_rag_documents(file_path):
    global _last_uploaded_pdf_path
    _last_uploaded_pdf_path = file_path

    # a new resume replaces the previous one -- clear out the old
    # TEMPORARY chunks (from both SQLite and the shared Chroma collection)
    # before ingesting, so the single collection doesn't accumulate every
    # past upload forever. PERMANENT chunks (explicitly remembered via
    # remember_info_tool) are untouched.
    removed = delete_temporary_chunks()
    if removed:
        print(f"[chroma] cleared {removed} temporary chunk(s) from the previous resume")

    # DB_PATH = './chroma_db'
    loader=PyPDFLoader(file_path)
    docs = loader.load()
    print(docs)

    strategy_name, spliter = pick_chunker_strategy(docs)
    print(f"[chunker] {len(docs)} page(s) -> using '{strategy_name}' strategy "
          f"(chunk_size={spliter._chunk_size}, overlap={spliter._chunk_overlap})")

    chunks = spliter.split_documents(docs)

    # full chunk text -> SQLite; Chroma embeds/stores only a short preview
    # per chunk (smaller index). chunk_id is stamped into each preview
    # Document's metadata so any retrieval path (which only has public
    # Chroma APIs available -- similarity_search etc. don't return raw ids)
    # can still join back to the full text in SQLite. Tagged "temporary" so
    # the next upload's delete_temporary_chunks() call clears these back out.
    chunk_ids = store_chunks(chunks, scope="temporary")
    preview_docs = [
        Document(
            page_content=chunk.page_content[:_CHROMA_PREVIEW_CHARS],
            metadata={**chunk.metadata, "chunk_id": chunk_id, "scope": "temporary"},
        )
        for chunk, chunk_id in zip(chunks, chunk_ids)
    ]

    # single shared collection ("my_pdf_document") for everything --
    # temporary and permanent chunks alike, distinguished only by the
    # "scope" metadata/SQLite column, instead of creating/dropping a
    # separate Chroma collection per upload.
    vector_store = Chroma.from_documents(
        documents = preview_docs,
        embedding = get_embedding_model(),
        collection_name='my_pdf_document',
        persist_directory='./my_chroma_db',
        ids = chunk_ids,
    )


def _resolve_full_text(doc):
    """Given a Document returned from Chroma (holding only a short preview
    plus a chunk_id in its metadata), returns its full text from SQLite --
    falling back to the preview if the chunk_id is missing/not found (e.g.
    documents ingested before this split existed)."""
    chunk_id = doc.metadata.get("chunk_id")
    if not chunk_id:
        return doc.page_content
    full_texts = load_chunk_texts([chunk_id])
    return full_texts.get(chunk_id, doc.page_content)


def get_resume_context(queries, k_per_query=3):
    """Retrieve just the resume chunks relevant to a set of section-focused
    queries (skills, experience, education, etc.) via the same
    Chroma/embedding pipeline ingest_rag_documents() already builds --
    reusing chunking+embeddings instead of feeding the small local LLM the
    entire raw resume text, which was overflowing its effective context and
    causing empty replies. Chroma only holds short previews, so full chunk
    text is looked up from SQLite via each Document's chunk_id. Chunks are
    deduped and returned in resume page order so the excerpt still reads
    coherently."""
    vector_store = Chroma(
        embedding_function=get_embedding_model(),
        collection_name='my_pdf_document',
        persist_directory='./my_chroma_db',
    )

    seen = {}
    for query in queries:
        for doc in vector_store.similarity_search(query, k=k_per_query):
            key = doc.metadata.get("chunk_id") or (doc.metadata.get("page"), doc.page_content[:50])
            seen[key] = doc

    ordered = sorted(seen.values(), key=lambda d: d.metadata.get("page", 0))
    return "\n\n".join(_resolve_full_text(doc) for doc in ordered)


def set_last_uploaded_jd(file_path):
    """Records the most recently uploaded job-description file path,
    used by resume_review_tool when the user attaches a JD instead of
    pasting it as text."""
    global _last_uploaded_jd_path
    _last_uploaded_jd_path = file_path

# -------------retriving data-------------
def get_retriever():
    vector_store=Chroma(
        embedding_function= get_embedding_model(),
        collection_name='my_pdf_document',
        persist_directory='./my_chroma_db'
    )
    retriver = vector_store.as_retriever(search_type='similarity', search_kwargs={'k': 4})
    return retriver
############## ---tools #####################


# --------------------------------Rag Tools-----------------
@tool
def remember_info_tool(info: str) -> str:
    """Use this when the user explicitly asks you to remember, save, or
    keep some piece of information permanently -- e.g. "remember that I
    prefer Python over Java", "save this for later: my target salary is
    X", "keep in mind that I'm looking for remote roles only". Stores the
    text as PERMANENT info in the shared document collection: unlike an
    uploaded resume (which is TEMPORARY and gets cleared out the next time
    a new resume is uploaded), permanent info stays available across future
    uploads and is included whenever rag_tool/resume review look things up.
    Do not use this for the resume/JD files themselves -- those are handled
    by the file upload flow, not this tool."""
    doc = Document(page_content=info, metadata={"source": "user_remembered_info"})
    strategy_name, spliter = pick_chunker_strategy([doc])
    chunks = spliter.split_documents([doc])

    chunk_ids = store_chunks(chunks, scope="permanent", id_prefix=f"{uuid.uuid4()}::")
    preview_docs = [
        Document(
            page_content=chunk.page_content[:_CHROMA_PREVIEW_CHARS],
            metadata={**chunk.metadata, "chunk_id": chunk_id, "scope": "permanent"},
        )
        for chunk, chunk_id in zip(chunks, chunk_ids)
    ]

    Chroma.from_documents(
        documents=preview_docs,
        embedding=get_embedding_model(),
        collection_name='my_pdf_document',
        persist_directory='./my_chroma_db',
        ids=chunk_ids,
    )

    return f"Remembered: \"{info}\" -- this will stay available across future uploads."


@tool
def rag_tool(query:str) -> str:
    """ Retrieve relavant infromation from pdf documents
    use this tool when user ask the factual concept questions that may be answered using the stored PDF documents """
    retriver = get_retriever()
    documents = retriver.invoke(query)
    if not documents:
        return "No relavant information found in the PDF"

    formatted_documents = []
    for index, document in enumerate(documents, start=1):
        source = document.metadata.get("source","unknown source")
        page = document.metadata.get("page","unknown source")

        formatted_documents.append(
            f"Document: {index}\n"
            f"Source: {source}\n"
            f"Page: {page} \n"
            f"Content: {_resolve_full_text(document)}"
        )
    return "\n\n".join(formatted_documents)


# --------------------------- ATS resume scoring tool -----------------------
@tool
def ats_score_tool(_: str = "") -> str:
    """Use this when the user asks to check/score/analyze their uploaded
    resume for ATS (Applicant Tracking System) compatibility, or asks things
    like "check ats score of this resume" / "fetch ats score" / "how ATS
    friendly is my resume". Returns targeted excerpts covering every resume
    section (contact/summary, skills, experience, projects, education), so
    the model can evaluate the whole document -- sections present,
    formatting, keyword coverage, contact info, etc. -- without being
    handed the entire raw PDF text, which overflows the local LLM's
    effective context and causes slow or empty replies."""
    report_progress("ats_score_tool", stage_index=0)  # Loading uploaded resume

    if not _last_uploaded_pdf_path:
        return (
            "No resume has been uploaded yet. Ask the user to upload their "
            "resume PDF using the file attachment in the chat box first."
        )

    report_progress("ats_score_tool", stage_index=1)  # Retrieving relevant resume sections

    # same targeted-chunk retrieval as resume_review_tool -- keeps the
    # payload small enough for a fast, reliable response instead of
    # dumping the full raw resume text
    resume_text = get_resume_context([
        "contact information summary professional profile",
        "skills technical skills core competencies",
        "work experience professional experience achievements",
        "projects",
        "education certifications",
    ])
    if not resume_text:
        resume_text = extract_full_text(_last_uploaded_pdf_path)

    report_progress("ats_score_tool", stage_index=2)  # Scoring ATS compatibility with the LLM

    return f"RESUME EXCERPTS (source: {_last_uploaded_pdf_path}):\n\n{resume_text}"


# --------------------------- Resume tailoring / review tool -----------------
RESUME_REVIEW_FOCUS_PROMPTS = {
    "recruiter_breakdown": (
        "Review the resume as if you're evaluating hundreds of candidates for a "
        "highly competitive role in the target industry. Give a raw, realistic "
        "assessment of how it performs in a real hiring environment. Point out "
        "every weak area that could lower shortlist chances -- unclear "
        "positioning, weak achievements, poor structure, generic wording, "
        "missing credibility signals, anything that makes the candidate look "
        "average compared to stronger candidates. State what instantly creates "
        "a negative impression, what feels forgettable, and what changes would "
        "make the resume look significantly more valuable, skilled, and "
        "interview-worthy."
    ),
    "attention_test": (
        "Evaluate the resume the way a recruiter would during the first 10 "
        "seconds of scanning it. State what stands out immediately, what feels "
        "forgettable, and whether it communicates enough value fast enough to "
        "compete with stronger applicants. Analyze layout, positioning, "
        "clarity, and first impression. Explain whether it makes the candidate "
        "look experienced, credible, and worth interviewing -- or just average."
    ),
    "ats_keywords": (
        "Compare the resume against the provided job description. Identify the "
        "technical terms, role-specific phrases, competencies, and industry "
        "keywords that are currently underrepresented or completely absent. "
        "Show how to naturally integrate those elements into the resume "
        "without making it feel stuffed or artificial. Recommend improvements "
        "to the experience section, skills layout, and wording strategy so the "
        "resume aligns more closely with this role and performs better during "
        "automated ATS screening."
    ),
    "impact_rewrite": (
        "Transform the experience section into high-value achievement "
        "statements that sound compelling to recruiters and hiring managers. "
        "Replace weak or responsibility-focused lines with results-oriented "
        "descriptions that clearly communicate ownership, contribution, and "
        "business impact. Wherever the resume lacks measurable outcomes, ask "
        "strategic follow-up questions to uncover useful metrics, "
        "improvements, efficiencies, revenue impact, time savings, or "
        "performance gains. Every bullet point should feel intentional, "
        "credible, and strong enough to justify an interview."
    ),
    "executive_polish": (
        "Conduct a deep final-pass evaluation from the perspective of a "
        "recruiter reviewing top-tier candidates. Identify anything that "
        "weakens clarity, professionalism, credibility, or overall impact -- "
        "inconsistent writing style, repetitive sentence structure, vague "
        "descriptions, weak transitions, low-value wording, overused "
        "corporate phrases that fail to differentiate the candidate. Refine "
        "the language throughout so every section sounds precise, "
        "intentional, and high-value. Replace generic statements with sharper "
        "positioning, stronger achievement-driven phrasing, and wording that "
        "reflects expertise, confidence, and real-world contribution. The "
        "final resume should feel polished, modern, and competitive."
    ),
    "market_positioning": (
        "Reshape the resume summary and core skills section to reflect the "
        "mindset, priorities, and terminology valued by the target industry "
        "and the companies the candidate is aiming for. Make the profile feel "
        "aligned with the industry's standards and culture instead of broad, "
        "outdated, or interchangeable with every other applicant -- "
        "specialized, relevant, and naturally suited for this type of role."
    ),
}


@tool
def resume_review_tool(
    job_description: str = "",
    focus_mode: str = "general",
) -> str:
    """Use this when the user wants their uploaded resume reviewed, critiqued,
    rewritten, or tailored to a specific job -- e.g. "review my resume like a
    recruiter", "give me a tailored resume for this JD", "rewrite my
    experience section with stronger impact", "optimize my resume for ATS
    keywords against this job description", "polish my resume", "align my
    resume to [company/industry]".

    Args:
        job_description: The target job description text, if the user pasted
            or mentioned one. Leave empty if none was given in the message --
            the tool will fall back to a job description file the user
            uploaded (if any).
        focus_mode: One of "recruiter_breakdown", "attention_test",
            "ats_keywords", "impact_rewrite", "executive_polish",
            "market_positioning", or "general" (runs a blended review
            covering all angles). Pick "ats_keywords" whenever a job
            description is available and the user wants tailoring; pick the
            others when the user's wording matches that specific angle
            (e.g. "first impression"/"10 seconds" -> attention_test,
            "rewrite my bullet points"/"make my experience sound stronger"
            -> impact_rewrite). Default to "general" if unsure.

    Before producing the final tailored resume text, this tool pauses and
    asks the human whether to add a specific ATS-boosting statement/keyword
    insertion it identified -- the caller (the LLM) should present that
    approval outcome to the user as part of the final answer.
    """
    report_progress("resume_review_tool", stage_index=0)  # Loading uploaded resume

    if not _last_uploaded_pdf_path:
        return (
            "No resume has been uploaded yet. Ask the user to upload their "
            "resume file (PDF, DOCX, or TXT) using the chat box's file "
            "attachment first."
        )

    report_progress("resume_review_tool", stage_index=1)  # Retrieving relevant resume sections

    # retrieve targeted chunks covering every resume section, via the
    # existing Chroma/embedding pipeline, instead of dumping the full raw
    # text -- keeps the payload small enough for the local LLM to actually
    # respond to, while still covering the whole document
    resume_text = get_resume_context([
        "contact information summary professional profile",
        "skills technical skills core competencies",
        "work experience professional experience achievements",
        "projects",
        "education certifications",
    ])
    if not resume_text:
        # retriever found nothing (e.g. ingestion hasn't finished/failed) --
        # fall back to the full extracted text rather than returning empty
        resume_text = extract_full_text(_last_uploaded_pdf_path)

    jd_text = job_description.strip()
    if not jd_text and _last_uploaded_jd_path:
        jd_text = extract_full_text(_last_uploaded_jd_path)

    focus_instructions = RESUME_REVIEW_FOCUS_PROMPTS.get(
        focus_mode,
        "Give a well-rounded review covering recruiter first impression, "
        "structural/wording weaknesses, achievement impact, and (if a job "
        "description is available) ATS keyword alignment.",
    )

    # HITL checkpoint: before finalizing, ask the human whether to add a
    # concrete ATS-boosting statement/keyword this tool has identified as a
    # likely score improvement, rather than silently inserting it. Dispatched
    # through the dynamic HITL step registry (see run_hitl_step above) so the
    # next step (offering a docx/pdf export) is chained in by the "ats_statement"
    # handler itself, not hardcoded here.
    suggested_statement = (
        "Proficient in applying [top missing keyword/skill from the job "
        "description] to deliver measurable results, demonstrated through "
        "[a specific quantifiable achievement from the resume]."
        if jd_text else
        "Delivered measurable, quantifiable impact through [specific "
        "achievement], directly supporting [team/business outcome]."
    )

    report_progress("resume_review_tool", stage_index=2)  # Waiting for your ATS-statement approval

    hitl_result = run_hitl_step(
        "ats_statement",
        question=(
            "To improve the ATS score, I'd like to add this statement to "
            "your resume (with the bracketed parts filled in from your "
            "real experience):\n\n"
            f"“{suggested_statement}”\n\n"
            "Add this statement to the tailored resume?"
        ),
        context={"suggested_statement": suggested_statement, "offer_export": True},
    )

    statement_note = hitl_result.get("message", "")
    export_note = (
        f"\n\nEXPORT: {hitl_result['message']}"
        if hitl_result.get("status") == "export_choice"
        else ""
    )

    jd_section = (
        f"JOB DESCRIPTION:\n{jd_text}\n\n"
        if jd_text else
        "JOB DESCRIPTION: (none provided -- skip JD-specific keyword "
        "matching and focus on general resume quality instead)\n\n"
    )

    report_progress("resume_review_tool", stage_index=4)  # Generating tailored resume with the LLM

    return (
        f"FOCUS: {focus_mode}\n\n"
        f"REVIEW INSTRUCTIONS:\n{focus_instructions}\n\n"
        f"HUMAN APPROVAL DECISION:\n{statement_note}{export_note}\n\n"
        f"{jd_section}"
        f"FULL RESUME TEXT (source: {_last_uploaded_pdf_path}):\n{resume_text}\n\n"
        f"After writing the tailored resume text above your answer, end your "
        f"reply with a line exactly formatted as:\n"
        f"===FINAL_TAILORED_RESUME_START===\n"
        f"<the complete final tailored resume text, ready to save as a file>\n"
        f"===FINAL_TAILORED_RESUME_END===\n"
        f"so the app can offer it as a downloadable file."
    )


# --------------------- ATS score -> gap analysis -> tailoring pipeline ----
@tool
def ats_tailor_pipeline_tool(job_description: str = "") -> str:
    """Use this for the FULL "tailor my resume for this job" request --
    e.g. "tailor my resume for this JD", "make my resume ATS-friendly for
    this job", "fix my resume so it passes ATS for this role", "what
    should I add to my resume for this job". This runs the complete
    internal pipeline in one call instead of the LLM having to chain
    separate tools itself:

      1. Load the uploaded resume.
      2. Load the job description (pasted text, an uploaded JD file, or if
         neither exists, the LLM should call search_tool FIRST to find a
         real JD and pass its text here as `job_description` -- same
         fallback rule as resume_review_tool).
      3. Score the resume's CURRENT ATS compatibility against that JD.
      4. Identify the specific missing keywords/skills and structural
         weaknesses an ATS system would flag.
      5. (Human-in-the-loop) confirm a suggested ATS-boosting statement.
      6. Produce the full resume back, annotated with WHERE each missing
         keyword should be inserted, plus the final tailored, ATS-
         optimized resume text.

    Args:
        job_description: The target job description text. Leave empty if
            the user didn't paste one -- the tool falls back to an
            uploaded JD file if present; if neither exists, the calling
            LLM should search for one first (see above) rather than call
            this with an empty JD.

    The final answer you give the user MUST include, in this order: (1)
    the ATS score and what's currently missing, (2) the full original
    resume with explicit notes on where each missing keyword/phrase
    should go, (3) the fully tailored resume text. This tool pauses
    mid-run for human approval, same as resume_review_tool -- present
    that outcome to the user as part of the answer.
    """
    report_progress("ats_tailor_pipeline_tool", stage_index=0)  # Loading uploaded resume

    if not _last_uploaded_pdf_path:
        return (
            "No resume has been uploaded yet. Ask the user to upload their "
            "resume file (PDF, DOCX, or TXT) using the chat box's file "
            "attachment first."
        )

    # ---- stage 1: load resume (full text -- the pipeline needs the WHOLE
    # document to point at exact insertion locations, not just top-k
    # similarity chunks like ats_score_tool/resume_review_tool use) ----
    resume_text = extract_full_text(_last_uploaded_pdf_path)

    report_progress("ats_tailor_pipeline_tool", stage_index=1)  # Loading/finding the job description

    jd_text = job_description.strip()
    if not jd_text and _last_uploaded_jd_path:
        jd_text = extract_full_text(_last_uploaded_jd_path)
    if not jd_text:
        return (
            "No job description is available yet. Search the web for a "
            "real, current job posting matching the role the user "
            "mentioned (via search_tool), then call ats_tailor_pipeline_tool "
            "again passing that job description text as `job_description` "
            "-- do not ask the user to paste one unless the search genuinely "
            "finds nothing usable."
        )

    report_progress("ats_tailor_pipeline_tool", stage_index=2)  # Scoring current ATS compatibility
    report_progress("ats_tailor_pipeline_tool", stage_index=3)  # Finding missing keywords and weaknesses

    suggested_statement = (
        "Proficient in applying [top missing keyword/skill from the job "
        "description] to deliver measurable results, demonstrated through "
        "[a specific quantifiable achievement from the resume]."
    )

    report_progress("ats_tailor_pipeline_tool", stage_index=4)  # Waiting for your ATS-statement approval

    hitl_result = run_hitl_step(
        "ats_statement",
        question=(
            "To improve the ATS score, I'd like to add this statement to "
            "your resume (with the bracketed parts filled in from your "
            "real experience):\n\n"
            f"“{suggested_statement}”\n\n"
            "Add this statement to the tailored resume?"
        ),
        context={"suggested_statement": suggested_statement, "offer_export": True},
    )

    statement_note = hitl_result.get("message", "")
    export_note = (
        f"\n\nEXPORT: {hitl_result['message']}"
        if hitl_result.get("status") == "export_choice"
        else ""
    )

    report_progress("ats_tailor_pipeline_tool", stage_index=6)  # Generating the tailored, ATS-optimized resume

    return (
        "ATS TAILORING PIPELINE -- follow these steps in your final answer, in order:\n\n"
        "STEP 1 -- CURRENT ATS SCORE:\n"
        "Compare the resume below against the job description below. Give a "
        "numeric or qualitative ATS compatibility score, and list the "
        "specific keywords/skills from the job description that are "
        "MISSING from the resume, plus any structural/formatting weaknesses "
        "an automated ATS parser would flag (e.g. missing standard section "
        "headings, inconsistent dates, no measurable achievements).\n\n"
        "STEP 2 -- FULL RESUME WITH INSERTION POINTS:\n"
        "Reproduce the user's full resume text below, and for each missing "
        "keyword/phrase identified in step 1, explicitly note WHERE in the "
        "resume it should be inserted (which section, near which existing "
        "bullet point) and a short suggested phrasing -- do not just say "
        "'add more keywords', point to the exact spot.\n\n"
        "STEP 3 -- FINAL TAILORED RESUME:\n"
        "Produce the complete, rewritten resume with all the missing "
        "keywords/phrases from step 2 actually incorporated naturally, "
        "ready to use as-is. Wrap ONLY this final tailored resume text in:\n"
        "===FINAL_TAILORED_RESUME_START===\n"
        "<the complete final tailored resume text>\n"
        "===FINAL_TAILORED_RESUME_END===\n"
        "so the app can offer it as a downloadable file.\n\n"
        f"HUMAN APPROVAL DECISION (mention this in your answer):\n{statement_note}{export_note}\n\n"
        f"JOB DESCRIPTION:\n{jd_text}\n\n"
        f"FULL RESUME TEXT (source: {_last_uploaded_pdf_path}):\n{resume_text}"
    )


# ---------------------------------------------------------------
#tools by default it's tool no need decorator sign
search_tool = TavilySearch(
    max_results = 5,
    topic='general',
    search_depth="advanced"
)
# ---------------------------------
#  calculator tools
@tool
def calculator(expression: str) -> str:
    """Evaluate a math expression, e.g. "2+2", "math.sqrt(16)", "10*5"."""

    try:
        allowed ={
            "math": math,
            "abs": abs,
            "round": round,
            "min": min,
            "max": max,
            "sum":sum
        }

        result= eval(expression, {"__builtins__":{}}, allowed)
        return str(result)
    except Exception as e:
        return f"Calculation error: {str(e)}"
# ---------------------------------------------------------------
#   stock tools 
def get_stock_price(symbol: str)-> dict:
    """
    fetch latest price for a given symbol (e.g. 'AAPL', 'TLSA')
    using Alpha vantage with API key in the URL
    """

    url =f"https://www.alphavantage.co/query?function=GLOBAL_QUOTE&symbol={symbol}&apikey={STOCK_API_KEY}"
    r= requests.get(url)
    return r.json()  

# ----------------------------------------------------------

@tool 
def purchase_tool(symbol:str, quantity:int) -> dict:
    """ Simulate purchase a given quantity of a stock symbol
    NOTE: tHIS IS MOCK IMPLEMENTATION 
    - NO REAL BREAKAGE API CALLED
    -IT SIMPLY RETURN CONFIRMATION PAYLOAD

    HUMAN IN THE LOOP  BEFORE CONFIRMING THE PURCHASE THIS TOOL WILL INTERRUPT AND WAIT FOR 
    FOR HUMAN DECISION {'YES'  / ANYTHING ELSE}
    """
    decision = interrupt(f"Approve buying {quantity} shares of symbol {symbol}?yes/no")
    if isinstance(decision, str) and decision.lower() in ("yes", "approved"):
        return {
             "status": "success",
        "message":f"purchase order placed for {quantity} shares of symbol {symbol}",
        "symbol":symbol,
        "quantity": quantity
        }
    else:

        return {
            "status": "cancelled",
            "message":f"purchase order cancelled for {quantity} shares of symbol {symbol}",
            "symbol":symbol,
            "quantity": quantity
        }
# -------------------------- weather tool ---------------
@tool
def get_weather(location: str) -> dict:
    """Get current weather (temp, humidity, wind) for a city, e.g. "Hyderabad, India"."""

    # 1. Geocode location -> latitude/longitude
    try:
        geo_response = requests.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={
                "name": location,
                "count": 1,
                "language": "en",
                "format": "json",
            },
            timeout=10,
        )
        geo_response.raise_for_status()
    except requests.exceptions.RequestException as e:
        # Open-Meteo's free public API has its own separate rate limit --
        # unrelated to the chat model's -- and can occasionally return 429
        # under bursty traffic. Report it gracefully instead of letting the
        # tool call crash.
        return {
            "success": False,
            "error": f"Weather service is temporarily unavailable ({e}). Please try again shortly.",
        }

    geo_data = geo_response.json()

    if not geo_data.get("results"):
        return {
            "success": False,
            "error": f"Location not found: {location}",
        }

    place = geo_data["results"][0]

    latitude = place["latitude"]
    longitude = place["longitude"]

    # 2. Get current weather
    try:
        weather_response = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": latitude,
                "longitude": longitude,
                "current": ",".join([
                    "temperature_2m",
                    "relative_humidity_2m",
                    "apparent_temperature",
                    "precipitation",
                    "weather_code",
                    "wind_speed_10m",
                ]),
                "timezone": "auto",
            },
            timeout=10,
        )
        weather_response.raise_for_status()
    except requests.exceptions.RequestException as e:
        return {
            "success": False,
            "error": f"Weather service is temporarily unavailable ({e}). Please try again shortly.",
        }

    weather_data = weather_response.json()

    return {
        "success": True,
        "location": {
            "name": place["name"],
            "country": place.get("country"),
            "latitude": latitude,
            "longitude": longitude,
        },
        "timezone": weather_data.get("timezone"),
        "current": weather_data.get("current"),
        "units": weather_data.get("current_units"),
    }
# ------------------------------------------------------------
# ------------------------- binding tool with bind_tools---------
# make tool list 

tools = [get_stock_price, search_tool, calculator, get_weather,rag_tool, purchase_tool, ats_score_tool, resume_review_tool, remember_info_tool, ats_tailor_pipeline_tool]

# make llm tool aware
llm_with_tools= llm.bind_tools(tools)
# -----------------------------------------------------------------

##################################################
# creating state

class ChatState(TypedDict):
    messages: Annotated[list[BaseMessage], add_messages]


STYLE_PROMPT = SystemMessage(content=(
    "For explain/describe/write-about answers: use Markdown (headings, "
    "bullets, numbered steps, code blocks, tables) to stay scannable, not "
    "one paragraph. For greetings or single-fact answers (calc, weather, "
    "stock price): reply in one plain line, no formatting."
))


def chat_node(state:ChatState):
    system_message=SystemMessage(content="""
You are a general-purpose AI assistant with tools. Pick the minimum tools
needed, never fabricate results, and answer directly from tool output
(don't just describe what a tool returned). Don't expose your tool-selection
reasoning.

For any request about an uploaded resume (scoring, reviewing, tailoring),
ALWAYS call the matching tool immediately -- never assume no resume was
uploaded and ask the user to upload one yourself. The tool itself checks
and will tell you if nothing has been uploaded; only relay that message to
the user if the tool actually reports it.

TOOLS:

- get_stock_price: current stock/market price.
- purchase_tool: buy stock; pauses for human approval. Reply based on the
  result's "status" field ("success"=placed, "cancelled"=rejected, NOT
  placed) -- never trust "message" text alone.
- search_tool: web search for current/external info not in an uploaded doc.
- calculator: any arithmetic/math -- never compute mentally.
- get_weather: current/forecast weather for a location.
- rag_tool: answer a specific factual question from an uploaded document
  using a few matched passages. NOT for whole-document evaluation (use
  ats_score_tool for that) -- rag_tool only sees snippets, not the full doc.
  If retrieved context is insufficient, say so; don't invent an answer.
- ats_score_tool: score/evaluate an uploaded resume's ATS compatibility
  (structure, keywords, sections). If none uploaded, ask the user to
  upload one -- don't guess a score.
- resume_review_tool(job_description, focus_mode): review/critique/rewrite
  a resume for narrower requests NOT covering the full score+tailor flow
  (e.g. "review my resume", "polish my language", a single focus angle).
  focus_mode: ats_keywords | attention_test | impact_rewrite |
  executive_polish | market_positioning | recruiter_breakdown | general.
  Pauses for human approval of a suggested ATS statement -- report that
  decision to the user.
- remember_info_tool(info): ONLY when the user explicitly says
  remember/save/keep something -- stores it PERMANENTLY (survives future
  resume uploads), unlike resume/JD uploads which are temporary.
- ats_tailor_pipeline_tool(job_description): PREFER over resume_review_tool
  for the FULL "tailor my resume for this job" ask (score -> find missing
  keywords -> human approval -> full resume with exact insertion points ->
  final tailored resume, all in one call). Your answer must cover, in
  order: (1) current ATS score + what's missing, (2) full original resume
  annotated with where to add each missing keyword, (3) final tailored
  resume wrapped in ===FINAL_TAILORED_RESUME_START/END===.

JD FALLBACK (resume_review_tool and ats_tailor_pipeline_tool): if the user
pasted a JD, pass it as job_description. If they uploaded a JD file
instead, leave job_description empty (auto-falls back to that file). If
NEITHER exists, call search_tool first to find a real JD for the role,
then call the tool again with that text -- never just ask the user to
paste one unless search finds nothing.

RAG: prefer rag_tool over search_tool for questions about an uploaded
document; combine both only if the user also wants current/external info
compared against it.
    """)
    # take user query from state
    messages = state['messages']
    #send to llm, with the tool-routing prompt + style instruction prepended
    try:
        response = llm_with_tools.invoke([system_message, STYLE_PROMPT] + messages)
    except Exception as e:
        # a raw API error (e.g. Groq's 413 "tokens per minute" rate limit
        # when a long resume/JD + conversation history pushes a request
        # over the free-tier budget) previously propagated uncaught,
        # silently failing the turn and leaving the LLM's next attempt
        # with no real tool result to work from -- producing a confusing
        # generic reply instead of an explanation. Surface it clearly.
        error_text = str(e)
        if "tokens per minute" in error_text.lower() or "rate_limit" in error_text.lower():
            # this fires on ANY request once the account's shared per-minute
            # token budget is used up -- including short/simple messages --
            # not only large resume/JD requests, so the explanation must not
            # imply the user's own message was too big.
            explanation = (
                "This app's AI model has hit its shared rate limit (tokens "
                "processed per minute) -- this isn't about the size of your "
                "message specifically; the limit is shared across all "
                "recent activity and needs a short time to reset. Please "
                "wait about a minute and try again."
            )
        else:
            explanation = f"The AI model backend returned an error: {error_text}"
        return {'messages': [AIMessage(content=explanation)]}
    #response to store state
    return {'messages': [response]}

tool_node = ToolNode(tools) # execute tool calls



# graph

conn= sqlite3.connect(database='chatbot.db', check_same_thread=False)
# checkpoint= MemorySaver()
checkpoint =SqliteSaver(conn)

# ----------------------draw nodes--------------
graph= StateGraph(ChatState)
graph.add_node('chatnode', chat_node)
graph.add_node("tools",tool_node)

# ---------------------add edges-------------
# add edges 
graph.add_edge(START, "chatnode")
#
#if the llm asked the tool, go to toolnode or else finish
graph.add_conditional_edges("chatnode", tools_condition)

graph.add_edge("tools","chatnode")


# ----------compile --------------
chatbot = graph.compile(checkpointer = checkpoint)

# -------------------threads operations------------
def get_all_threads():
    all_threads= set()
    for ckpt in checkpoint.list(None):
        all_threads.add(ckpt.config['configurable']['thread_id'])
    return list(all_threads)


def delete_thread(thread_id):
    """Permanently remove a thread's checkpoints and pending writes."""
    conn.execute("DELETE FROM checkpoints WHERE thread_id = ?", (thread_id,))
    conn.execute("DELETE FROM writes WHERE thread_id = ?", (thread_id,))
    conn.commit()
# ----------------------------------------------

# thread_id = "1"
# initial_state = {
#     "messages" :[
#         HumanMessage(content= "What is my name")
#     ]
# }
# config={'configurable':{'thread_id': thread_id}}
# response = chatbot.invoke(initial_state, config=config)
# print(response['messages'][-1].content)
    
